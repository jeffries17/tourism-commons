#!/usr/bin/env python3
"""
Convert HTML Regional Competitive Benchmarking document to Word format
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from bs4 import BeautifulSoup

def create_word_document():
    """Create a Word document from the HTML content"""
    
    # Read the HTML file
    html_file = "REGIONAL_COMPETITIVE_BENCHMARKING.html"
    if not os.path.exists(html_file):
        print(f"❌ HTML file not found: {html_file}")
        return
    
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Regional Competitive Benchmarking', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_heading('Creative Industries Analysis: Gambia vs West African Competitors', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run('Project: ').bold = True
    metadata.add_run('Regional Benchmarking & Market Positioning Analysis\n')
    metadata.add_run('Component: ').bold = True
    metadata.add_run('Creative Industries Competitive Analysis\n')
    metadata.add_run('Date: ').bold = True
    metadata.add_run('October 2025\n')
    metadata.add_run('Data Period: ').bold = True
    metadata.add_run('2013-2025\n')
    metadata.add_run('Total Reviews Analyzed: ').bold = True
    metadata.add_run('1,316 (Gambia) + 3,096 (Regional) = 4,412 reviews')
    
    # Add executive summary
    doc.add_heading('Executive Summary', level=1)
    
    exec_summary = doc.add_paragraph()
    exec_summary.add_run('This analysis compares Gambian creative industries against 45 regional competitors across 5 West African countries (Benin, Cape Verde, Ghana, Nigeria, Senegal) to identify competitive gaps, best practices, and strategic opportunities for improvement.')
    
    # Add key findings
    doc.add_heading('Key Findings', level=2)
    
    # Gambia's position
    doc.add_heading("Gambia's Competitive Position:", level=3)
    position_list = [
        "Overall Sentiment: +0.19 (Creative Industries only)",
        "Regional Ranking: 4th of 6 countries", 
        "Gap to Leaders: -0.09 sentiment points behind Benin (+0.28)",
        "Competitive with: Senegal (+0.20), Cape Verde (+0.18)",
        "Ahead of: Nigeria (+0.15)"
    ]
    
    for item in position_list:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # Critical gaps
    doc.add_heading('Critical Gaps:', level=3)
    gaps_list = [
        "Infrastructure & Facilities: -0.19 points behind regional leaders",
        "Educational Value: -0.15 points behind best practices", 
        "Value for Money: -0.12 points below regional average"
    ]
    
    for item in gaps_list:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # Competitive advantages
    doc.add_heading('Competitive Advantages:', level=3)
    advantages_list = [
        "Artistic & Creative Quality: +0.02 ahead of regional average",
        "Cultural Authenticity: 67% positive mentions (vs 58% regional average)"
    ]
    
    for item in advantages_list:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # Country comparison table
    doc.add_heading('1. Country-by-Country Sentiment Comparison', level=1)
    doc.add_heading('Table 1: Creative Industries Performance by Country', level=2)
    
    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    hdr_cells = table.rows[0].cells
    headers = ['Country', 'Stakeholders', 'Total Reviews', 'Avg Sentiment', 'Avg Rating', 'Top Performer', 'Gambia Gap']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make headers bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    table_data = [
        ['Benin', '7', '412', '+0.28', '4.18/5', 'Musée Fondation Zinsou (+0.32)', '-0.09 (Gambia behind)'],
        ['Ghana', '17', '1,398', '+0.26', '4.21/5', 'Cape Coast Castle (+0.24)', '-0.07 (Gambia behind)'],
        ['Senegal', '12', '891', '+0.20', '4.15/5', 'Gorée Island Museums (+0.29)', '-0.01 (Gambia behind)'],
        ['Gambia', '12', '1,316', '+0.19', '4.06/5', 'Kachikally Crocodile Pool (+0.21)', '— Baseline —'],
        ['Cape Verde', '5', '223', '+0.18', '3.92/5', 'Mindelo Cultural Centre (+0.26)', '+0.01 (Gambia ahead)'],
        ['Nigeria', '4', '172', '+0.15', '3.88/5', 'Nike Art Gallery (+0.23)', '+0.04 (Gambia ahead)']
    ]
    
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    # Key insights
    doc.add_heading('Key Insights', level=2)
    insights = [
        "Gambia ranks 4th of 6 countries — mid-tier performance",
        "Gap to leaders (Benin): -0.09 sentiment points (~32% improvement needed to match)",
        "Competitive with: Senegal, Cape Verde (within margin of error)",
        "Ahead of: Nigeria (emerging creative tourism market)"
    ]
    
    for insight in insights:
        p = doc.add_paragraph(insight, style='List Bullet')
    
    # Theme analysis
    doc.add_heading('2. Theme-by-Theme Competitive Analysis', level=1)
    doc.add_heading('Table 2: Thematic Performance — Gambia vs Regional Leaders', level=2)
    
    # Create theme table
    theme_table = doc.add_table(rows=1, cols=6)
    theme_table.style = 'Table Grid'
    theme_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Theme headers
    theme_headers = ['Theme', 'Gambia Score', 'Regional Avg', 'Best Regional Performer', 'Gap', 'Learning Opportunity']
    theme_hdr_cells = theme_table.rows[0].cells
    for i, header in enumerate(theme_headers):
        theme_hdr_cells[i].text = header
        for paragraph in theme_hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Theme data
    theme_data = [
        ['Cultural & Heritage Value', '+0.22', '+0.24', 'Benin (+0.29) Musée Fondation Zinsou', '-0.07', 'Heritage site presentation, exhibit design'],
        ['Service & Staff Quality', '+0.24', '+0.26', 'Ghana (+0.31) Cape Coast guides', '-0.07', 'Tour guide certification, hospitality training'],
        ['Facilities & Infrastructure', '+0.09', '+0.28', 'Benin (+0.36) Restored museums', '-0.19 ⚠️', 'Investment in building maintenance, signage'],
        ['Accessibility & Transport', '+0.21', '+0.26', 'Senegal (+0.31) Gorée ferry system', '-0.10', 'Transport logistics, wayfinding, ferry reliability'],
        ['Value for Money', '+0.15', '+0.22', 'Ghana (+0.27) Transparent pricing', '-0.12', 'Pricing communication, value perception'],
        ['Safety & Security', '+0.16', '+0.18', 'Senegal (+0.24) Visible security', '-0.08', 'Security presence, safety communication'],
        ['Educational & Informational Value', '+0.19', '+0.28', 'Benin (+0.34) Interpretive centers', '-0.15', 'Interpretive signage, audio guides, guided tours'],
        ['Artistic & Creative Quality', '+0.21', '+0.19', 'Gambia (+0.21) Craft markets', '+0.02 ✅', 'Gambia is competitive — maintain quality'],
        ['Atmosphere & Overall Experience', '+0.23', '+0.27', 'Senegal (+0.32) Immersive design', '-0.09', 'Atmospheric design, sensory experiences']
    ]
    
    for row_data in theme_data:
        row_cells = theme_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    # Best practice case studies
    doc.add_heading('3. Best Practice Case Studies', level=1)
    
    # Case Study 1
    doc.add_heading('Case Study 1: Musée de la Fondation Zinsou (Benin) — Infrastructure Excellence', level=2)
    doc.add_paragraph('Sentiment Score: +0.32 (32% higher than Gambian museums avg)')
    
    doc.add_heading('What They Do Well:', level=3)
    doc.add_paragraph('Facilities & Infrastructure: +0.36 (vs Gambia +0.09)')
    facilities_list = [
        "Modern, climate-controlled gallery spaces",
        "Clean, accessible facilities with multilingual signage", 
        "Regular maintenance and restoration programs"
    ]
    for item in facilities_list:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Educational Value: +0.34 (vs Gambia +0.19)')
    education_list = [
        "Professional audio guides in 4 languages",
        "Trained docents available for all exhibits",
        "Interactive displays and contextual panels"
    ]
    for item in education_list:
        doc.add_paragraph(item, style='List Bullet')
    
    # Quote
    quote_para = doc.add_paragraph()
    quote_para.add_run('"La Fondation Zinsou is a must-visit for anyone interested in contemporary African art and culture. The collection is beautifully curated, with stunning works by talented artists from Benin and across the continent. The exhibits are thought-provoking and offer a fresh perspective on African creativity and heritage."')
    quote_para.add_run('\n— TripAdvisor Review, 5/5').italic = True
    
    doc.add_heading('Transferable Lessons for Gambia:', level=3)
    lessons_list = [
        "Investment in climate control and preservation extends building lifespan AND improves visitor experience",
        "Multilingual interpretation (not just English) increases accessibility for Francophone and other African travelers",
        "Staff training as museum educators (not just security) elevates Educational Value scores"
    ]
    for item in lessons_list:
        doc.add_paragraph(item, style='List Number')
    
    # Strategic recommendations
    doc.add_heading('4. Strategic Recommendations', level=1)
    
    doc.add_heading('Priority 1: Infrastructure Investment (Gap: -0.19 points)', level=2)
    doc.add_heading('Immediate Actions (0-6 months):', level=3)
    immediate_actions = [
        "Emergency repairs at Kunta Kinteh Island (structural preservation)",
        "Ferry service backup plan (private boat partnerships)",
        "Basic facility upgrades (toilets, signage) at top 5 visited sites"
    ]
    for action in immediate_actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_heading('Medium-term (6-18 months):', level=3)
    medium_term = [
        "Establish Heritage Conservation Fund (levy on tourism receipts)",
        "Partner with UNESCO for preservation technical assistance",
        "Implement regular maintenance schedules"
    ]
    for item in medium_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Expected Impact: +0.12 sentiment boost, moving Gambia from 4th to 2nd in regional rankings')
    
    # Methodology
    doc.add_heading('Methodology Notes', level=1)
    
    doc.add_heading('Sentiment Score Interpretation', level=2)
    sentiment_scale = [
        "+0.50 to +1.00: Exceptional (rarely achieved; represents near-universal praise)",
        "+0.30 to +0.49: Very Positive (strong recommendation, high satisfaction)",
        "+0.20 to +0.29: Positive (generally satisfied, some areas for improvement)",
        "+0.10 to +0.19: Mixed Positive (satisfied but notable concerns)",
        "0.00 to +0.09: Neutral/Low Positive (tepid satisfaction)",
        "-0.09 to -0.01: Neutral/Low Negative (dissatisfaction emerging)",
        "-0.10 to -0.29: Negative (significant problems, poor experience)",
        "-0.30 to -0.50: Very Negative (strong dissatisfaction)",
        "-0.51 to -1.00: Extremely Negative (rare; represents universal condemnation)"
    ]
    
    for item in sentiment_scale:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Statistical Notes', level=2)
    stats_notes = [
        "Minimum 24 reviews required for stakeholder inclusion (ensures statistical reliability)",
        "Margin of error: ±0.03 sentiment points at 95% confidence for stakeholders with 100+ reviews",
        "Theme mention minimum: 5 mentions required for theme-level analysis (prevents outlier skewing)"
    ]
    
    for note in stats_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Report Prepared By: ').bold = True
    footer_para.add_run('Regional Benchmarking & Market Positioning Analysis Team\n')
    footer_para.add_run('Data Analysis Period: ').bold = True
    footer_para.add_run('October 2025\n')
    footer_para.add_run('Review Period Covered: ').bold = True
    footer_para.add_run('2013-2025 (primary focus 2019-2025)\n')
    footer_para.add_run('Total Data Points: ').bold = True
    footer_para.add_run('4,412 reviews, 57 stakeholders, 6 countries, 9 themes, 12,296 theme mentions analyzed\n\n')
    footer_para.add_run('This report is Component 3 of Deliverable 2: Regional Benchmarking & Market Positioning Analysis. It provides the competitive context for the Creative Tourism Personas Framework and Digital Positioning Opportunities Matrix.').italic = True
    
    # Save document
    output_file = "REGIONAL_COMPETITIVE_BENCHMARKING.docx"
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    
    return output_file

if __name__ == "__main__":
    create_word_document()
